#!/usr/bin/env python3
"""
MHP Construction Projects - Document Analysis Script
Analyzes .docx files to extract scope of work patterns and pricing references.
"""

import os
import sys
import json
import re
from collections import defaultdict

# Install python-docx if needed
try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    from docx import Document


def get_all_docx_files(base_path):
    """Walk directory tree and collect all .docx files."""
    all_files = []
    for root, dirs, files in os.walk(base_path):
        for f in files:
            if f.endswith('.docx') and not f.startswith('~$'):
                all_files.append(os.path.join(root, f))
    return all_files


def categorize_files(all_files):
    """Categorize files by document type based on filename patterns."""
    categories = defaultdict(list)
    for fp in all_files:
        fn = os.path.basename(fp).lower()
        if any(x in fn for x in ['scope of service', 'scope of servcies', '(scope of services)', 'scope of services']):
            categories['scope_of_services'].append(fp)
        elif any(x in fn for x in ['project estimate', 'project proposal', 'estimate binder']):
            categories['estimates_proposals'].append(fp)
        elif any(x in fn for x in ['change order', 'overage memo', 'allowance']):
            categories['change_orders_memos'].append(fp)
        elif 'contract' in fn:
            categories['contracts'].append(fp)
        elif any(x in fn for x in ['project report', 'progress report', 'closing project', 'closeout']):
            categories['project_reports'].append(fp)
        elif 'intake' in fn:
            categories['intake'].append(fp)
        else:
            categories['other'].append(fp)
    return dict(categories)


def read_docx_text(filepath):
    """Read all text from a docx file."""
    try:
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also read tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    tables_text.append(' | '.join(row_text))

        return {
            'paragraphs': paragraphs,
            'tables': tables_text,
            'full_text': '\n'.join(paragraphs + tables_text)
        }
    except Exception as e:
        return {'error': str(e), 'paragraphs': [], 'tables': [], 'full_text': ''}


def extract_project_type(filename, folder_path, text):
    """Determine the project type from filename, folder, and text content."""
    combined = (filename + ' ' + folder_path + ' ' + text[:500]).lower()

    project_types = []
    type_keywords = {
        'porch': ['porch', 'front porch', 'back porch', 'screened porch'],
        'deck': ['deck', 'decking'],
        'addition': ['addition', 'room addition'],
        'kitchen_renovation': ['kitchen', 'kitchen reno', 'kitchen remodel'],
        'bathroom_renovation': ['bathroom', 'bath reno', 'bath remodel', 'shower'],
        'roofing': ['roof', 'roofing', 'shingle', 're-roof'],
        'siding': ['siding', 'hardie', 'vinyl siding'],
        'garage_carport': ['garage', 'carport'],
        'guest_house_adu': ['guest house', 'adu', 'accessory dwelling'],
        'foundation_repair': ['foundation', 'foundation repair'],
        'retaining_wall': ['retaining wall'],
        'concrete': ['concrete', 'driveway', 'sidewalk'],
        'painting': ['paint', 'painting', 'exterior paint', 'interior paint'],
        'door_replacement': ['door replacement', 'door install', 'entry door', 'french door'],
        'window_replacement': ['window', 'window replacement'],
        'flooring': ['flooring', 'hardwood', 'lvp', 'tile floor'],
        'fencing': ['fence', 'fencing'],
        'commercial_buildout': ['vet clinic', 'veterinary', 'commercial', 'office buildout', 'tenant'],
        'water_damage_repair': ['water damage', 'moisture', 'mold'],
        'exterior_renovation': ['exterior', 'exterior renovation', 'curb appeal'],
        'gutter': ['gutter', 'gutters'],
        'electrical': ['electrical', 'wiring', 'panel'],
        'plumbing': ['plumbing', 'pipe', 'water heater'],
        'hvac': ['hvac', 'heating', 'cooling', 'air condition', 'ductwork'],
        'insulation': ['insulation', 'spray foam'],
        'demolition': ['demo', 'demolition', 'tear down'],
        'framing': ['framing', 'structural'],
        'landscaping': ['landscape', 'landscaping', 'yard', 'grading'],
        'staircase': ['stair', 'staircase', 'steps'],
    }

    for ptype, keywords in type_keywords.items():
        for kw in keywords:
            if kw in combined:
                project_types.append(ptype)
                break

    return list(set(project_types)) if project_types else ['general_construction']


def extract_scope_items(text):
    """Extract scope/line items from document text."""
    items = []
    lines = text.split('\n')

    for line in lines:
        line = line.strip()
        # Look for numbered items, bulleted items, or items with pricing
        if re.match(r'^[\d]+[\.\)]\s+', line) or \
           re.match(r'^[-•]\s+', line) or \
           re.match(r'^[A-Z][a-z].*:', line) or \
           ('install' in line.lower() or 'remove' in line.lower() or
            'repair' in line.lower() or 'replace' in line.lower() or
            'build' in line.lower() or 'construct' in line.lower() or
            'demolish' in line.lower() or 'demo' in line.lower() or
            'paint' in line.lower() or 'frame' in line.lower()):
            if len(line) > 10 and len(line) < 500:
                items.append(line)

    return items


def extract_pricing(text):
    """Extract pricing references from text."""
    prices = []
    # Match dollar amounts
    price_patterns = re.findall(r'\$[\d,]+(?:\.\d{2})?', text)
    for p in price_patterns:
        # Find context around the price
        idx = text.find(p)
        start = max(0, idx - 80)
        end = min(len(text), idx + len(p) + 40)
        context = text[start:end].strip()
        prices.append({'amount': p, 'context': context})

    return prices


def extract_change_order_info(text, filename):
    """Extract change order specific information."""
    info = {
        'filename': os.path.basename(filename),
        'prices': extract_pricing(text),
        'items': [],
        'category': 'unknown'
    }

    # Determine change order category
    fn_lower = filename.lower()
    text_lower = text.lower()

    co_categories = {
        'flooring': ['flooring', 'floor', 'hardwood', 'lvp', 'tile'],
        'allowance_overage': ['allowance', 'overage'],
        'cabinets_countertops': ['cabinet', 'countertop', 'granite'],
        'electrical': ['electrical', 'wiring', 'outlet', 'switch'],
        'plumbing': ['plumbing', 'pipe', 'fixture'],
        'structural': ['structural', 'beam', 'joist', 'foundation'],
        'doors_windows': ['door', 'window'],
        'painting': ['paint', 'stain'],
        'gutters_rails': ['gutter', 'rail', 'railing'],
        'shower': ['shower'],
        'roofing': ['roof', 'shingle'],
        'siding': ['siding'],
        'concrete': ['concrete'],
        'framing': ['framing', 'frame'],
        'hvac': ['hvac', 'heating', 'cooling'],
        'insulation': ['insulation'],
        'demolition': ['demo', 'demolition'],
        'design_change': ['design change', 'plan change', 'revision'],
        'material_upgrade': ['upgrade', 'material change'],
        'additional_scope': ['additional', 'add-on', 'added scope'],
    }

    for cat, keywords in co_categories.items():
        for kw in keywords:
            if kw in fn_lower or kw in text_lower[:500]:
                info['category'] = cat
                break
        if info['category'] != 'unknown':
            break

    # Extract line items from change order
    info['items'] = extract_scope_items(text)

    return info


def main():
    base_path = '/Users/jameswalton/Desktop/Construction Projects'
    output_path = '/Users/jameswalton/Desktop/MHPEstimate/data/scope_analysis.json'

    print("Step 1: Collecting all .docx files...")
    all_files = get_all_docx_files(base_path)
    print(f"  Found {len(all_files)} .docx files")

    print("\nStep 2: Categorizing files...")
    categories = categorize_files(all_files)
    for cat, files in sorted(categories.items()):
        print(f"  {cat}: {len(files)} files")

    # ---- ANALYZE SCOPE OF SERVICES ----
    print("\nStep 3: Analyzing Scope of Services documents...")
    scope_docs = categories.get('scope_of_services', [])
    sample_scope = scope_docs[:15]  # Take up to 15

    project_types_catalog = defaultdict(lambda: {
        'count': 0,
        'example_projects': [],
        'standard_scope_items': [],
        'sample_documents': []
    })

    all_scope_items_by_type = defaultdict(list)

    for fp in sample_scope:
        print(f"  Reading: {os.path.basename(fp)}")
        data = read_docx_text(fp)
        if data.get('error'):
            print(f"    Error: {data['error']}")
            continue

        folder = os.path.dirname(fp)
        ptypes = extract_project_type(os.path.basename(fp), folder, data['full_text'])
        items = extract_scope_items(data['full_text'])

        for pt in ptypes:
            project_types_catalog[pt]['count'] += 1
            project_types_catalog[pt]['example_projects'].append(os.path.basename(fp))
            project_types_catalog[pt]['sample_documents'].append(os.path.basename(fp))
            all_scope_items_by_type[pt].extend(items)

    # ---- ANALYZE ESTIMATES/PROPOSALS ----
    print("\nStep 4: Analyzing Estimates/Proposals...")
    estimate_docs = categories.get('estimates_proposals', [])
    sample_estimates = estimate_docs[:15]

    pricing_data = []

    for fp in sample_estimates:
        print(f"  Reading: {os.path.basename(fp)}")
        data = read_docx_text(fp)
        if data.get('error'):
            print(f"    Error: {data['error']}")
            continue

        folder = os.path.dirname(fp)
        ptypes = extract_project_type(os.path.basename(fp), folder, data['full_text'])
        prices = extract_pricing(data['full_text'])
        items = extract_scope_items(data['full_text'])

        estimate_info = {
            'document': os.path.basename(fp),
            'project_types': ptypes,
            'prices': prices,
            'line_items': items[:30],  # limit
            'table_data': data['tables'][:20]
        }
        pricing_data.append(estimate_info)

        for pt in ptypes:
            project_types_catalog[pt]['count'] += 1
            if os.path.basename(fp) not in project_types_catalog[pt]['example_projects']:
                project_types_catalog[pt]['example_projects'].append(os.path.basename(fp))
            all_scope_items_by_type[pt].extend(items)

    # ---- ANALYZE CHANGE ORDERS ----
    print("\nStep 5: Analyzing Change Order Memos...")
    co_docs = categories.get('change_orders_memos', [])
    sample_cos = co_docs[:15]

    change_order_data = []
    co_categories_count = defaultdict(int)

    for fp in sample_cos:
        print(f"  Reading: {os.path.basename(fp)}")
        data = read_docx_text(fp)
        if data.get('error'):
            print(f"    Error: {data['error']}")
            continue

        co_info = extract_change_order_info(data['full_text'], fp)
        co_info['table_data'] = data['tables'][:15]
        change_order_data.append(co_info)
        co_categories_count[co_info['category']] += 1

    # ---- BUILD FINAL ANALYSIS ----
    print("\nStep 6: Building final analysis...")

    # Deduplicate and limit scope items per type
    for pt in all_scope_items_by_type:
        items = list(set(all_scope_items_by_type[pt]))
        project_types_catalog[pt]['standard_scope_items'] = items[:50]

    # Convert defaultdict to regular dict for JSON serialization
    project_types_final = {}
    for pt, info in project_types_catalog.items():
        project_types_final[pt] = {
            'count': info['count'],
            'example_projects': list(set(info['example_projects']))[:10],
            'standard_scope_items': info['standard_scope_items'],
            'sample_documents': list(set(info['sample_documents']))[:5]
        }

    # Also scan ALL filenames to get broader project type distribution
    print("\nStep 7: Scanning all filenames for project type distribution...")
    all_project_types_count = defaultdict(int)
    all_project_folders = set()

    for fp in all_files:
        folder_name = fp.replace(base_path, '').split('/')[1] if '/' in fp.replace(base_path, '') else ''
        if folder_name:
            all_project_folders.add(folder_name)
        ptypes = extract_project_type(os.path.basename(fp), os.path.dirname(fp), '')
        for pt in ptypes:
            all_project_types_count[pt] += 1

    # Build the output
    analysis = {
        'metadata': {
            'total_documents_analyzed': len(all_files),
            'scope_documents_sampled': len(sample_scope),
            'estimate_documents_sampled': len(sample_estimates),
            'change_order_documents_sampled': len(sample_cos),
            'analysis_date': '2026-03-07',
            'source_directory': base_path
        },
        'document_categories': {
            cat: {
                'count': len(files),
                'sample_filenames': [os.path.basename(f) for f in files[:8]]
            }
            for cat, files in sorted(categories.items())
        },
        'project_folders': sorted(list(all_project_folders)),
        'project_types_catalog': project_types_final,
        'project_type_distribution_all_files': dict(sorted(all_project_types_count.items(), key=lambda x: -x[1])),
        'pricing_references': pricing_data,
        'change_order_patterns': {
            'documents_analyzed': len(change_order_data),
            'category_distribution': dict(co_categories_count),
            'details': change_order_data,
            'common_change_order_categories': sorted(co_categories_count.keys(), key=lambda x: -co_categories_count[x])
        },
        'common_line_items_by_project_category': {
            pt: items[:30] for pt, items in all_scope_items_by_type.items()
        }
    }

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    with open(output_path, 'w') as f:
        json.dump(analysis, f, indent=2)

    print(f"\nAnalysis complete! Written to: {output_path}")
    print(f"Project types found: {list(project_types_final.keys())}")
    print(f"Change order categories: {list(co_categories_count.keys())}")


if __name__ == '__main__':
    main()
